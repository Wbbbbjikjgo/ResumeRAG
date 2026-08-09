"""Document loader with format routing."""

from pathlib import Path
from typing import BinaryIO

from loguru import logger

from app.core.config import get_settings


class DocumentLoader:
    """Load documents from various formats."""

    @staticmethod
    def detect_format(filename: str, content_type: str | None = None) -> str:
        """Detect document format from filename and content type."""
        ext = Path(filename).suffix.lower()
        
        format_map = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "doc",
            ".html": "html",
            ".htm": "html",
            ".txt": "text",
            ".png": "image",
            ".jpg": "image",
            ".jpeg": "image",
        }
        
        fmt = format_map.get(ext)
        if fmt:
            return fmt
        
        # Fallback to content type
        if content_type:
            if "pdf" in content_type:
                return "pdf"
            elif "word" in content_type or "docx" in content_type:
                return "docx"
            elif "html" in content_type:
                return "html"
            elif "image" in content_type:
                return "image"
        
        return "unknown"

    @staticmethod
    async def load(file_content: bytes, filename: str, content_type: str | None = None) -> dict:
        """
        Load document and extract text.
        
        Returns:
            dict with keys: text, format, metadata
        """
        fmt = DocumentLoader.detect_format(filename, content_type)
        logger.info(f"Loading document: {filename} (format: {fmt})")
        
        if fmt == "pdf":
            return await DocumentLoader._load_pdf(file_content, filename)
        elif fmt == "docx":
            return await DocumentLoader._load_docx(file_content, filename)
        elif fmt == "html":
            return await DocumentLoader._load_html(file_content, filename)
        elif fmt == "text":
            return await DocumentLoader._load_text(file_content, filename)
        elif fmt == "image":
            # Image will be handled by OCR module
            return {
                "text": "",
                "format": "image",
                "metadata": {"filename": filename, "needs_ocr": True}
            }
        else:
            logger.warning(f"Unknown format for {filename}, attempting text extraction")
            return await DocumentLoader._load_text(file_content, filename)

    @staticmethod
    async def _load_pdf(content: bytes, filename: str) -> dict:
        """Load PDF document."""
        import pdfplumber
        import io
        
        text_parts = []
        metadata = {"filename": filename, "pages": 0, "needs_ocr": False}
        
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            metadata["pages"] = len(pdf.pages)
            
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                else:
                    # Page has no text, might be scanned image
                    metadata["needs_ocr"] = True
        
        text = "\n\n".join(text_parts)
        
        # If most pages need OCR, flag the whole document
        if metadata["needs_ocr"] and len(text_parts) < metadata["pages"] * 0.5:
            metadata["needs_ocr"] = True
        else:
            metadata["needs_ocr"] = False
        
        logger.info(f"PDF loaded: {filename}, {metadata['pages']} pages, text length: {len(text)}")
        
        return {
            "text": text,
            "format": "pdf",
            "metadata": metadata
        }

    @staticmethod
    async def _load_docx(content: bytes, filename: str) -> dict:
        """Load Word document."""
        from docx import Document
        import io
        
        doc = Document(io.BytesIO(content))
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        text = "\n".join(text_parts)
        
        metadata = {
            "filename": filename,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables)
        }
        
        logger.info(f"DOCX loaded: {filename}, {metadata['paragraphs']} paragraphs, text length: {len(text)}")
        
        return {
            "text": text,
            "format": "docx",
            "metadata": metadata
        }

    @staticmethod
    async def _load_html(content: bytes, filename: str) -> dict:
        """Load HTML document."""
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(content, "html.parser")
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text(separator="\n", strip=True)
        
        metadata = {"filename": filename}
        logger.info(f"HTML loaded: {filename}, text length: {len(text)}")
        
        return {
            "text": text,
            "format": "html",
            "metadata": metadata
        }

    @staticmethod
    async def _load_text(content: bytes, filename: str) -> dict:
        """Load plain text document."""
        text = content.decode("utf-8", errors="ignore")
        
        metadata = {"filename": filename}
        logger.info(f"Text loaded: {filename}, text length: {len(text)}")
        
        return {
            "text": text,
            "format": "text",
            "metadata": metadata
        }
